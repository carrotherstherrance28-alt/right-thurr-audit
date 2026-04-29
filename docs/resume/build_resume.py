from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


INK = "1A1714"
ORANGE = "D9621F"
MUTED = "574E45"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="E5DCCD"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        borders.append(tag)
    tc_pr.append(borders)


def add_bottom_border(paragraph, color=ORANGE, size="12"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p_pr.append(borders)


def set_run(run, *, size=None, bold=None, color=None, font="Arial"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def para(doc, text="", style=None, spacing_after=3, spacing_before=0):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.space_before = Pt(spacing_before)
    if text:
        r = p.add_run(text)
        set_run(r, size=9.5, color=INK)
    return p


def section_heading(doc, text):
    p = para(doc, spacing_after=5, spacing_before=9)
    r = p.add_run(text.upper())
    set_run(r, size=9, bold=True, color=ORANGE)
    add_bottom_border(p)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.paragraph_format.space_after = Pt(2.5)
    r = p.add_run(text)
    set_run(r, size=9.2, color=INK)
    return p


def role(doc, title, org, location, dates, description=None):
    p = para(doc, spacing_after=1, spacing_before=5)
    r = p.add_run(f"{title} | {org} | {location} | {dates}")
    set_run(r, size=10, bold=True, color=INK)
    if description:
        d = para(doc, description, spacing_after=3)
        d.runs[0].italic = True
        d.runs[0].font.color.rgb = RGBColor.from_string(MUTED)


def make_resume():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(9.2)

    name = para(doc, spacing_after=0)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = name.add_run("Therrance Carrothers")
    set_run(r, size=20, bold=True, color=INK, font="Georgia")

    title = para(doc, spacing_after=2)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("AI Automation Specialist | Workflow Builder | Brand-Aware Systems Designer")
    set_run(r, size=9.5, bold=True, color=ORANGE)

    contact = para(
        doc,
        "Houston, TX | Open to relocation | (314) 817-7538 | carrothers.therrance28@gmail.com | LinkedIn | Portfolio",
        spacing_after=6,
    )
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.runs[0].font.size = Pt(8.5)
    contact.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    section_heading(doc, "Summary")
    para(
        doc,
        "AI Automation Specialist with hands-on experience designing and deploying end-to-end business systems using n8n, LLMs, APIs, Supabase, and workflow orchestration. Strong record improving lead throughput, reducing manual work, and building reliable automation layers around sales, onboarding, reporting, and operations. Adds brand-aware UI, design-system, and content production skills to make automation work visible, usable, and client-ready.",
        spacing_after=5,
    )

    section_heading(doc, "Core Skills")
    table = doc.add_table(rows=4, cols=2)
    table.autofit = False
    widths = [Inches(1.55), Inches(5.95)]
    rows = [
        ("AI & Automation", "n8n, Make.com, Zapier, OpenRouter, ChatGPT, Claude, LangChain, Cohere, 11Labs, Whisper, OCR tools"),
        ("Technical Stack", "APIs, JSON, Webhooks, Python, JavaScript, Postman, Firecrawl, RAG, Supabase, Qdrant, Pinecone, Retell AI"),
        ("Design & Content", "Brand systems, responsive landing pages, UI wireframes, Claude Design, v0.app, Lovable, Remotion video templates, content workflows"),
        ("Business & Ops", "Process mapping, ROI analysis, stakeholder interviews, prompt engineering, workflow documentation, Jira, Miro, Notion, Slack"),
    ]
    for row, (label, body) in zip(table.rows, rows):
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
            set_cell_border(row.cells[idx])
        set_cell_shading(row.cells[0], "F5EFE2")
        p_label = row.cells[0].paragraphs[0]
        p_body = row.cells[1].paragraphs[0]
        p_label.paragraph_format.space_after = Pt(0)
        p_body.paragraph_format.space_after = Pt(0)
        r_label = p_label.add_run(label)
        r_body = p_body.add_run(body)
        set_run(r_label, size=8.5, bold=True, color=ORANGE)
        set_run(r_body, size=8.5, color=INK)

    section_heading(doc, "Work Experience")
    role(
        doc,
        "AI Automation Specialist",
        "Maximax",
        "Houston, TX",
        "01/2021 - Present",
        "Automation work spanning an AI automation agency, marketing firm, UK-based accounting company, and construction consultancy.",
    )
    bullet(doc, "Designed and deployed a 29-node AI lead intelligence pipeline using n8n, OpenRouter, Tavily, and Bouncer.io, delivering 668% first-year ROI with a 16-day payback period.")
    bullet(doc, "Scaled lead throughput from 4 to 30/hour (+650%) while reducing cost per lead from $8.75 to $1.17 through AI enrichment and workflow automation.")
    bullet(doc, "Cut invalid outreach from 18% to under 2% using AI-powered email validation and LLM-based lead scoring, reducing wasted SDR time.")
    bullet(doc, "Reduced proposal turnaround from 24-48 hours to under 10 minutes by automating document generation and contract delivery.")
    bullet(doc, "Compressed client onboarding from 90 minutes to under 2 minutes by integrating n8n, Google Drive, and Airtable.")
    bullet(doc, "Built client-facing automation assets with clearer user flows, branded reporting language, and design-aware layouts for proposals, diagnostics, and handoffs.")
    bullet(doc, "Achieved 95%+ reliability across production workflows through systematic error handling and fallback logic design.")

    role(
        doc,
        "AI Operations Manager",
        "Ace Pursuit Trucking LLC",
        "Houston, TX",
        "08/2021 - Present",
        "Trucking and freight operation with $60K+/month logistics activity.",
    )
    bullet(doc, "Managed end-to-end operations across 20+ vendor accounts while maintaining a 98% on-time payment rate and reducing reconciliation errors by 35%.")
    bullet(doc, "Designed financial tracking systems that cut reporting time by 40% and improved month-end close accuracy to under 2% variance.")
    bullet(doc, "Eliminated redundant manual workflows, reducing coordination overhead by 30% and saving approximately 10 hours per month.")
    bullet(doc, "Converted operational and financial patterns into automation requirements, including process maps, variance controls, and owner visibility workflows.")

    section_heading(doc, "Selected System Outcomes")
    bullet(doc, "Lead intelligence engine: +650% throughput, lower cost per lead, stronger validation, and LLM-assisted prioritization.")
    bullet(doc, "Proposal automation: document generation and contract delivery reduced from days to minutes.")
    bullet(doc, "Client onboarding system: intake, folders, Airtable records, and handoff structure automated for repeatable setup.")
    bullet(doc, "Brand-aware buildout workflow: landing page, diagnostic report, Discord alert, Supabase intake, and Remotion content starter for Thurr Solutions.")

    section_heading(doc, "Education & Certifications")
    para(doc, "Bachelor in Business Administration (Finance), University of Missouri - Columbia, Trulaske College of Business", spacing_after=2)
    para(doc, "AI Automation Specialist, Careerist.com | AWS Certified Cloud Practitioner (in progress)", spacing_after=0)

    return doc


if __name__ == "__main__":
    output = "docs/resume/Therrance_Carrothers_AI_Automation_Resume.docx"
    make_resume().save(output)
    print(output)
